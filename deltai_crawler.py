#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DeltAI — volledige site-crawler voor deltawonen.nl (HTML + PDF + DOCX)
- Respecteert robots.txt
- Start op sitemap.xml + volgt interne links
- Beperkt tot deltawonen.nl
- Extraheert hoofdtekst uit HTML (readability-lxml)
- Extraheert tekst uit PDF (pdfminer.six) en DOCX (python-docx)
- Output:
  data/raw_pages.jsonl  (ruwe tekst + metadata per URL)
  data/qa.jsonl         (Q&A-items: vraag = titel/bestandsnaam, antwoord = samenvatting, + bron)
  data/assets/          (gedownloade PDF/DOCX)
Gebruik:
  python -m venv .venv
  # Windows: .venv\Scripts\activate
  # macOS/Linux: source .venv/bin/activate
  pip install -r requirements.txt
  python deltai_crawler.py --max-pages 3000 --delay 1.5
"""
import argparse, os, re, time, json, queue, urllib.parse as up
from dataclasses import dataclass, asdict
from typing import List, Set, Dict, Tuple
import requests
from bs4 import BeautifulSoup
from readability import Document
from lxml.html.clean import Cleaner
from urllib import robotparser
from tqdm import tqdm
from io import BytesIO
from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document as DocxDocument

ALLOWED_NETLOCS = {"deltawonen.nl", "www.deltawonen.nl"}
USER_AGENT = "DeltAI-Scraper/1.0"
DEFAULT_START = ["https://www.deltawonen.nl/"]
SITEMAPS = ["https://www.deltawonen.nl/sitemap.xml"]

DATA_DIR = "data"
RAW_JSONL = os.path.join(DATA_DIR, "raw_pages.jsonl")
QA_JSONL = os.path.join(DATA_DIR, "qa.jsonl")
ASSETS_DIR = os.path.join(DATA_DIR, "assets")
os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(ASSETS_DIR, exist_ok=True)

@dataclass
class PageItem:
    url: str
    content_type: str  # html|pdf|docx|other|error
    title: str
    text: str
    status: int
    fetched_at: str
    sha1: str
    links: List[str]

def sha1_text(s: str) -> str:
    import hashlib
    return hashlib.sha1(s.encode("utf-8","ignore")).hexdigest()

def is_allowed(url: str) -> bool:
    try:
        p = up.urlparse(url)
        if p.scheme not in {"http","https"}: return False
        if p.netloc.lower() not in ALLOWED_NETLOCS: return False
        if re.search(r"\.(jpg|jpeg|png|gif|svg|webp|ico|css|js|woff2?|ttf|eot|mp4|webm|zip|rar)(\?.*)?$", p.path, re.I):
            return False
        return True
    except:
        return False

def absolutize(base: str, href: str) -> str:
    return up.urljoin(base, href.strip())

def get_robot_parser(root: str):
    rp = robotparser.RobotFileParser()
    rp.set_url(up.urljoin(root, "/robots.txt"))
    try: rp.read()
    except: pass
    return rp

def get_sitemap_urls() -> List[str]:
    urls=[]
    for sm in SITEMAPS:
        try:
            r = requests.get(sm, headers={"User-Agent": USER_AGENT}, timeout=20)
            if r.status_code==200 and "xml" in r.headers.get("content-type",""):
                soup = BeautifulSoup(r.text, "xml")
                for loc in soup.find_all("loc"):
                    u = (loc.text or "").strip()
                    if is_allowed(u): urls.append(u)
        except:
            pass
    return list(dict.fromkeys(urls))

def fetch(url: str, timeout: int=30) -> Tuple[int,str,bytes,Dict[str,str]]:
    r = requests.get(url, headers={"User-Agent": USER_AGENT}, timeout=timeout)
    return r.status_code, r.headers.get("content-type",""), r.content, dict(r.headers)

def clean_html_to_text(html: str):
    doc = Document(html)
    title = (doc.short_title() or "").strip()
    article_html = doc.summary(html_partial=True)
    cleaner = Cleaner(style=True, scripts=True, comments=True, javascript=True,
                      page_structure=False, safe_attrs_only=False)
    cleaned = cleaner.clean_html(article_html)
    soup = BeautifulSoup(cleaned, "html5lib")
    for tag in soup.select("nav, footer, header, script, style"):
        tag.decompose()
    text = soup.get_text(" ", strip=True)
    return title, text

def extract_pdf_text(content: bytes) -> str:
    try:
        return pdf_extract_text(BytesIO(content)) or ""
    except Exception as e:
        return f"[[PDF extract failed: {e}]]"

def extract_docx_text(content: bytes) -> str:
    try:
        doc = DocxDocument(BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        return f"[[DOCX extract failed: {e}]]"

def guess_filename(url: str, headers: Dict[str,str]) -> str:
    import os
    name = os.path.basename(up.urlparse(url).path) or "download"
    cd = headers.get("Content-Disposition") or headers.get("content-disposition","")
    import re
    m = re.search(r'filename="?([^"]+)"?', cd)
    if m: name = m.group(1)
    return re.sub(r"[^A-Za-z0-9._-]+","_", name)

def discover_links(base_url: str, html: str) -> List[str]:
    soup = BeautifulSoup(html, "html.parser")
    out=[]
    for a in soup.find_all("a", href=True):
        u = absolutize(base_url, a["href"])
        if is_allowed(u): out.append(u)
    return list(dict.fromkeys(out))

def html_is_article(text: str) -> bool:
    return len(text.split()) >= 50

def write_jsonl(path: str, item: dict):
    with open(path, "a", encoding="utf-8") as f:
        f.write(json.dumps(item, ensure_ascii=False) + "\n")

def summarize(text: str, max_chars: int=1000) -> str:
    txt = " ".join((text or "").split())
    return txt[:max_chars] + ("…" if len(txt) > max_chars else "")

def main():
    import datetime
    ap = argparse.ArgumentParser()
    ap.add_argument("--start", nargs="*", default=DEFAULT_START)
    ap.add_argument("--max-pages", type=int, default=3000)
    ap.add_argument("--delay", type=float, default=1.5)
    ap.add_argument("--include-docs", action="store_true")
    args = ap.parse_args()

    rp = get_robot_parser(args.start[0])
    def allowed_by_robots(u):
        try: return rp.can_fetch(USER_AGENT, u)
        except: return True

    seen: Set[str] = set()
    q = queue.Queue()

    for u in (get_sitemap_urls() + list(args.start)):
        if is_allowed(u) and allowed_by_robots(u):
            q.put(u)

    pbar = tqdm(total=args.max_pages, desc="Crawling")
    count=0

    while not q.empty() and count < args.max_pages:
        url = q.get()
        if url in seen:
            continue
        seen.add(url)

        try:
            status, ctype, content, headers = fetch(url)
        except Exception as e:
            write_jsonl(RAW_JSONL, asdict(PageItem(
                url=url, content_type="error", title="", text=f"[[FETCH ERROR: {e}]]",
                status=0, fetched_at=datetime.datetime.utcnow().isoformat(), sha1="", links=[]
            )))
            continue

        fetched_at = datetime.datetime.utcnow().isoformat()
        ctype_low = (ctype or "").lower()
        links=[]; title=""; text=""

        if status==200:
            if "text/html" in ctype_low:
                html = content.decode("utf-8","ignore")
                title, text = clean_html_to_text(html)
                if not html_is_article(text):
                    soup = BeautifulSoup(html, "html.parser")
                    text = soup.get_text(" ", strip=True)

                links = discover_links(url, html)
                for u2 in links:
                    if u2 not in seen and is_allowed(u2) and allowed_by_robots(u2):
                        q.put(u2)

                write_jsonl(RAW_JSONL, asdict(PageItem(
                    url=url, content_type="html", title=title, text=text, status=status,
                    fetched_at=fetched_at, sha1=sha1_text(text), links=links
                )))
                if title and text:
                    write_jsonl(QA_JSONL, {
                        "question": title,
                        "answer": summarize(text, 1000),
                        "tags": ["deltawonen","html"],
                        "sources": [url]
                    })

            elif "application/pdf" in ctype_low or url.lower().endswith(".pdf"):
                fname = guess_filename(url, headers)
                with open(os.path.join(ASSETS_DIR, fname),"wb") as f: f.write(content)
                text = extract_pdf_text(content)
                write_jsonl(RAW_JSONL, asdict(PageItem(
                    url=url, content_type="pdf", title=fname, text=text, status=status,
                    fetched_at=fetched_at, sha1=sha1_text(text), links=[]
                )))
                write_jsonl(QA_JSONL, {
                    "question": fname.replace(".pdf","").replace("_"," ").strip(),
                    "answer": summarize(text, 1000),
                    "tags": ["deltawonen","pdf"],
                    "sources": [url]
                })

            elif "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in ctype_low or url.lower().endswith(".docx"):
                fname = guess_filename(url, headers)
                with open(os.path.join(ASSETS_DIR, fname),"wb") as f: f.write(content)
                text = extract_docx_text(content)
                write_jsonl(RAW_JSONL, asdict(PageItem(
                    url=url, content_type="docx", title=fname, text=text, status=status,
                    fetched_at=fetched_at, sha1=sha1_text(text), links=[]
                )))
                write_jsonl(QA_JSONL, {
                    "question": fname.replace(".docx","").replace("_"," ").strip(),
                    "answer": summarize(text, 1000),
                    "tags": ["deltawonen","docx"],
                    "sources": [url]
                })

            else:
                if args.include_docs and re.search(r"\.(doc|xls|ppt)(x)?$", url, re.I):
                    fname = guess_filename(url, headers)
                    with open(os.path.join(ASSETS_DIR, fname),"wb") as f: f.write(content)
                write_jsonl(RAW_JSONL, asdict(PageItem(
                    url=url, content_type=ctype_low, title="", text="", status=status,
                    fetched_at=fetched_at, sha1="", links=[]
                )))
        else:
            write_jsonl(RAW_JSONL, asdict(PageItem(
                url=url, content_type="error", title="", text=f"HTTP {status}", status=status,
                fetched_at=fetched_at, sha1="", links=[]
            )))

        count += 1
        pbar.update(1)
        time.sleep(args.delay)

    pbar.close()
    print(f"Done. Visited {count} pages.")
    print(f"Outputs:\n - {RAW_JSONL}\n - {QA_JSONL}\n - assets in {ASSETS_DIR}")

if __name__ == "__main__":
    main()
